from docx import Document
from docx.shared import Pt
from requests import get
from platform import system
from os import startfile
from subprocess import call
from string import ascii_uppercase

# Data Configs
kahootID = []
rawData = []
compData = []
docQue = Document()
docAns = Document()

# Copyright & Instructions
print('--------------------------------')
print('Copyright (c) 2020-2021 Hao Kang')
print('(* Always Hit ENTER to Continue)')
print('--------------------------------')
print('Copy and Paste Kahoot URL: ')

# Capture Multiple URLs
while True:
    URL = input()
    if URL != '':
        kahootID.append(URL.split('/')[-1])
    else:
        break

# Instructions
print('--------------------------------')
print('URL are Captured → Ready to Run')
print('--------------------------------\n')

# Kahoot API
for a in kahootID:
    answerURL = 'https://create.kahoot.it/rest/kahoots/{}/card/?includeKahoot=true'.format(a)
    rawData.extend(get(answerURL).json()['kahoot']['questions'])
    [compData.append(x) for x in rawData if x not in compData]


# Purify Text
def purify(original):
    original = original.replace('&nbsp;', '')
    original = original.replace('<i>', '')
    original = original.replace('</i>', '')
    original = original.replace('<b>', '')
    original = original.replace('</b>', '')
    return original


# Improved Paragraph Format
def paragraph_format(original):
    original.paragraph_format.line_spacing = 1.0
    original.paragraph_format.space_before = Pt(0)
    original.paragraph_format.space_after = Pt(8)


# Loop-i: Capture Questions
for i in range(0, len(compData)):
    text = purify('{ques}'.format(ques=compData[i]['question']))
    paragraphQue = docQue.add_paragraph(text, style='List Number')
    paragraph_format(paragraphQue)

    # Loop-j: Capture Choices
    for j in range(0, len(compData[i]['choices'])):
        text = purify('\t{ansNum}. {ans}'.format(ansNum=list(ascii_uppercase)[j], ans=compData[i]['choices'][j]['answer']))
        paragraphQue = docQue.add_paragraph(text)
        paragraph_format(paragraphQue)

    # Capture Answers
    ans = []
    [ans.append(k) for k in range(0, len(compData[i]['choices'])) if compData[i]['choices'][k]['correct']]
    text = ''
    for x in ans:
        text += list(ascii_uppercase)[x]
    paragraphAns = docAns.add_paragraph(text, style='List Number')

# Save Document
docQue.save('KahootExport-Question.docx')
docAns.save('KahootExport-Answer.docx')
print('---------------------------------')
print('Saved: KahootExport-Question.docx')
print('Saved: KahootExport-Answer.docx')
print('---------------------------------')

# Exit and Open
input('\n(* Hit ENTER to Exit and Open Documents)\n')

# Mac
if system() == 'Darwin':
    call(('open', 'KahootExport-Question.docx'))
    call(('open', 'KahootExport-Answer.docx'))
# Windows
elif system() == 'Windows':
    startfile('KahootExport-Question.docx')
    startfile('KahootExport-Answer.docx')
# Linux
else:
    call(('xdg-open', 'KahootExport-Question.docx'))
    call(('xdg-open', 'KahootExport-Answer.docx'))
